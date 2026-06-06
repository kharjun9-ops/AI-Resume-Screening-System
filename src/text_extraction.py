import PyPDF2
import docx
import os
from io import BytesIO

try:
    import fitz  # PyMuPDF (much faster PDF text extraction)
except Exception:
    fitz = None

try:
    from PIL import Image
    import pytesseract

    _TESSERACT_CMD = os.getenv("TESSERACT_CMD")
    if _TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = _TESSERACT_CMD
    else:
        for candidate in (
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
        ):
            if os.path.exists(candidate):
                pytesseract.pytesseract.tesseract_cmd = candidate
                break
except Exception:
    Image = None
    pytesseract = None

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    max_pages = int(os.getenv("PDF_MAX_PAGES", "0"))
    if fitz is not None:
        try:
            with fitz.open(pdf_path) as doc:
                page_count = doc.page_count
                limit = page_count if max_pages <= 0 else min(page_count, max_pages)
                return "\n".join(doc.load_page(i).get_text("text") for i in range(limit))
        except Exception as e:
            print(f"Error reading PDF file {pdf_path} with PyMuPDF: {e}")

    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_index, page in enumerate(pdf_reader.pages):
                if max_pages > 0 and page_index >= max_pages:
                    break
                extracted = page.extract_text() or ""
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error reading PDF file {pdf_path}: {e}")
    return text


def extract_text_from_pdf_bytes(pdf_bytes: bytes):
    """Extract text from PDF bytes (avoids saving uploads to disk)."""
    max_pages = int(os.getenv("PDF_MAX_PAGES", "0"))
    if not pdf_bytes:
        return ""

    if fitz is not None:
        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                page_count = doc.page_count
                limit = page_count if max_pages <= 0 else min(page_count, max_pages)
                return "\n".join(doc.load_page(i).get_text("text") for i in range(limit))
        except Exception as e:
            print(f"Error reading PDF bytes with PyMuPDF: {e}")

    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        parts = []
        for page_index, page in enumerate(pdf_reader.pages):
            if max_pages > 0 and page_index >= max_pages:
                break
            extracted = page.extract_text() or ""
            if extracted:
                parts.append(extracted)
        return "\n".join(parts)
    except Exception as e:
        print(f"Error reading PDF bytes: {e}")
        return ""


def extract_text_from_docx_bytes(docx_bytes: bytes):
    """Extract text from DOCX bytes (avoids saving uploads to disk)."""
    if not docx_bytes:
        return ""
    try:
        doc = docx.Document(BytesIO(docx_bytes))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        print(f"Error reading DOCX bytes: {e}")
        return ""


def extract_text_from_image_bytes(image_bytes: bytes):
    """Extract text from image bytes using OCR."""
    if Image is None or pytesseract is None:
        print("OCR dependencies not installed. Install Pillow and pytesseract.")
        return "OCR not available."

    if not image_bytes:
        return "OCR not available."

    try:
        with Image.open(BytesIO(image_bytes)) as img:
            return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Error reading image bytes: {e}")
        return "OCR not available."


def extract_text_from_bytes(file_bytes: bytes, filename: str):
    """Extracts text from bytes based on the filename extension."""
    _, file_extension = os.path.splitext(filename)
    file_extension = file_extension.lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    if file_extension == ".docx":
        return extract_text_from_docx_bytes(file_bytes)
    if file_extension == ".txt":
        try:
            return (file_bytes or b"").decode("utf-8", errors="ignore")
        except Exception:
            return ""
    if file_extension in IMAGE_EXTENSIONS:
        return extract_text_from_image_bytes(file_bytes)

    return "Unsupported file type."

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX file {docx_path}: {e}")
    return text

def extract_text_from_image(image_path):
    """Extracts text from an image file using Tesseract OCR."""
    if Image is None or pytesseract is None:
        print("OCR dependencies not installed. Install Pillow and pytesseract.")
        return "OCR not available."

    try:
        with Image.open(image_path) as img:
            return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Error reading image file {image_path}: {e}")
        return "OCR not available."

def extract_text_from_file(file_path):
    """Extracts text from a file based on its extension."""
    if not os.path.exists(file_path):
        return "File not found."

    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif file_extension in IMAGE_EXTENSIONS:
        return extract_text_from_image(file_path)
    else:
        return "Unsupported file type."

if __name__ == '__main__':
    # Example Usage
    # Create dummy files for testing
    if not os.path.exists('dummy.txt'):
        with open('dummy.txt', 'w') as f:
            f.write("This is a test text file.")

    # You will need to provide your own PDF and DOCX files for a full test
    print("Text from TXT:", extract_text_from_file('dummy.txt'))

    # To test with a PDF, you would do something like this:
    # print("Text from PDF:", extract_text_from_file('path/to/your/resume.pdf'))

    # To test with a DOCX, you would do something like this:
    # print("Text from DOCX:", extract_text_from_file('path/to/your/resume.docx'))
